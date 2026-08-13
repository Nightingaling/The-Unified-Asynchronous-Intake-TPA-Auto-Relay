import secrets
import time
from html import escape

import streamlit as st
from docx import Document


def read_docx_text(uploaded_file):
	uploaded_file.seek(0)
	document = Document(uploaded_file)
	text_blocks = [
		paragraph.text.strip()
		for paragraph in document.paragraphs
		if paragraph.text.strip()
	]

	for table in document.tables:
		for row in table.rows:
			row_text = " | ".join(
				cell.text.strip() for cell in row.cells if cell.text.strip()
			)
			if row_text:
				text_blocks.append(row_text)

	return "\n".join(text_blocks)


def render_copy_button(reference_code):
	st.iframe(
		f"""
		<style>
			body {{ margin: 0; font-family: Georgia, "Times New Roman", serif; }}
			button {{
				width: 100%;
				min-height: 58px;
				border: 2px solid #176da8;
				border-radius: 6px;
				background: #176da8;
				color: #ffffff;
				font: 800 17px Arial, sans-serif;
				cursor: pointer;
			}}
			button:hover, button:focus {{ background: #0f5688; border-color: #0f5688; }}
			button:focus-visible {{ outline: 3px solid #f2b84b; outline-offset: 2px; }}
			#copy-status {{
				min-height: 20px;
				margin-top: 4px;
				color: #277a58;
				font: 700 14px Arial, sans-serif;
				text-align: center;
			}}
		</style>
		<button id="copy-reference" type="button">Copy Reference Code</button>
		<div id="copy-status" role="status" aria-live="polite"></div>
		<script>
			const referenceCode = "{reference_code}";
			const button = document.getElementById("copy-reference");
			const status = document.getElementById("copy-status");

			button.addEventListener("click", async () => {{
				try {{
					if (navigator.clipboard && window.isSecureContext) {{
						await navigator.clipboard.writeText(referenceCode);
					}} else {{
						const textArea = document.createElement("textarea");
						textArea.value = referenceCode;
						textArea.style.position = "fixed";
						textArea.style.opacity = "0";
						document.body.appendChild(textArea);
						textArea.select();
						document.execCommand("copy");
						textArea.remove();
					}}

					button.textContent = "Copied!";
					status.textContent = "Reference code copied to clipboard.";
				}} catch (error) {{
					status.textContent = "Unable to copy. Please write down the code shown above.";
				}}
			}});
		</script>
		""",
		height=86,
	)


st.set_page_config(
	page_title="Pre-Arrival Registration | Medical Centre",
	page_icon="🏥",
	layout="centered",
	initial_sidebar_state="collapsed",
)

if "step" not in st.session_state:
	st.session_state.step = 0
if "reference_code" not in st.session_state:
	st.session_state.reference_code = None

st.markdown(
	"""
	<style>
		:root {
			--navy: #123453;
			--blue: #176da8;
			--blue-dark: #0f5688;
			--green: #277a58;
			--ink: #172b3a;
			--muted: #526674;
			--line: #cbd7df;
			--surface: #ffffff;
			--wash: #f3f8fa;
		}

		.stApp {
			background: linear-gradient(180deg, #e9f4f7 0, #f8fbfc 22rem, #ffffff 42rem);
			color: var(--ink);
		}

		[data-testid="stHeader"],
		[data-testid="stToolbar"] {
			display: none;
		}

		.main .block-container {
			max-width: 48rem;
			padding: 1.25rem 1.5rem 3rem;
		}

		.site-header {
			display: flex;
			align-items: center;
			justify-content: space-between;
			gap: 1rem;
			padding-bottom: 1rem;
			border-bottom: 2px solid var(--navy);
		}

		.brand {
			display: flex;
			align-items: center;
			gap: 0.75rem;
			color: var(--navy);
			font-family: Georgia, "Times New Roman", serif;
			font-size: 1.3rem;
			font-weight: 700;
		}

		.brand-mark {
			display: grid;
			width: 2.6rem;
			height: 2.6rem;
			border-radius: 50%;
			background: var(--navy);
			color: white;
			font-family: Arial, sans-serif;
			font-size: 1.8rem;
			place-items: center;
		}

		.help-number {
			color: var(--ink);
			font-size: 1rem;
			font-weight: 700;
			text-align: right;
		}

		.help-number span {
			display: block;
			color: var(--muted);
			font-size: 0.8rem;
			font-weight: 600;
			text-transform: uppercase;
		}

		.hero {
			padding: 4rem 0 2.5rem;
		}

		h1, h2, h3, p, label, button, input {
			letter-spacing: 0 !important;
		}

		.hero h1, .flow-title h1 {
			max-width: 42rem;
			margin: 0 0 1rem;
			color: var(--navy);
			font-family: Georgia, "Times New Roman", serif;
			font-size: 3rem;
			font-weight: 700;
			line-height: 1.12;
		}

		.hero p, .flow-title p {
			max-width: 40rem;
			margin: 0;
			color: var(--muted);
			font-size: 1.2rem;
			line-height: 1.65;
		}

		.assurance-row {
			display: grid;
			grid-template-columns: repeat(3, 1fr);
			gap: 0;
			margin: 2rem 0 0;
			border-top: 1px solid var(--line);
			border-bottom: 1px solid var(--line);
			background: rgba(255, 255, 255, 0.72);
		}

		.assurance-item {
			padding: 1.1rem 1rem;
			color: var(--ink);
			font-size: 1rem;
			font-weight: 700;
			text-align: center;
		}

		.assurance-item + .assurance-item {
			border-left: 1px solid var(--line);
		}

		.progress-label {
			margin: 2rem 0 0.45rem;
			color: var(--muted);
			font-size: 0.95rem;
			font-weight: 700;
		}

		.progress-track {
			height: 0.7rem;
			overflow: hidden;
			border-radius: 0.35rem;
			background: #d9e3e8;
		}

		.progress-fill {
			height: 100%;
			border-radius: inherit;
			background: var(--green);
		}

		.flow-title {
			padding: 2.5rem 0 1.5rem;
		}

		.flow-title h1 {
			font-size: 2.25rem;
		}

		.bot-message {
			margin: 0 0 1.5rem;
			padding: 1.25rem 1.35rem;
			border-left: 5px solid var(--blue);
			border-radius: 0 6px 6px 0;
			background: #eaf4fa;
			color: var(--ink);
			font-size: 1.1rem;
			line-height: 1.55;
		}

		[data-testid="stTextInput"] label,
		[data-testid="stFileUploader"] label {
			color: var(--ink);
			font-size: 1.1rem;
			font-weight: 700;
		}

		[data-testid="stTextInput"] input {
			min-height: 3.5rem;
			border: 2px solid #9fb1bd;
			border-radius: 6px;
			font-size: 1.1rem;
		}

		[data-testid="stFileUploader"] {
			padding: 0;
			border: 0;
			background: transparent;
		}

		[data-testid="stFileUploaderDropzone"] {
			display: flex;
			align-items: center;
			justify-content: center;
			flex-direction: column;
			gap: 1rem;
			min-height: 10rem;
			padding: 1.5rem !important;
			border: 2px dashed #6f91a6;
			border-radius: 6px;
			background: var(--wash);
			text-align: center;
		}

		[data-testid="stFileUploaderDropzoneInstructions"] {
			width: 100%;
			color: var(--ink);
			text-align: center;
		}

		[data-testid="stFileUploaderDropzone"] svg {
			width: 2rem;
			height: 2rem;
			color: var(--blue);
		}

		[data-testid="stFileUploaderDropzone"] button {
			position: static !important;
			width: auto;
			min-height: 3rem;
			padding: 0.65rem 1.5rem !important;
			border: 2px solid var(--blue) !important;
			border-radius: 6px !important;
			background: #ffffff !important;
			color: var(--blue) !important;
			font-size: 1rem;
			font-weight: 800;
		}

		[data-testid="stFileUploaderDropzone"] button:hover,
		[data-testid="stFileUploaderDropzone"] button:focus {
			background: #e6f2f8 !important;
			color: var(--blue-dark) !important;
		}

		.stButton > button, .stDownloadButton > button {
			width: 100%;
			min-height: 3.7rem;
			margin-top: 0.65rem;
			border: 2px solid var(--blue);
			border-radius: 6px;
			background: var(--blue);
			color: #ffffff;
			font-size: 1.1rem;
			font-weight: 800;
		}

		.stButton > button:hover, .stDownloadButton > button:hover {
			border-color: var(--blue-dark);
			background: var(--blue-dark);
			color: #ffffff;
		}

		[data-testid="stAlert"] {
			font-size: 1.05rem;
		}

		.completion {
			padding: 2.5rem 0 1rem;
			text-align: center;
		}

		.success-mark {
			display: grid;
			width: 4.5rem;
			height: 4.5rem;
			margin: 0 auto 1.25rem;
			border-radius: 50%;
			background: var(--green);
			color: white;
			font-size: 2.5rem;
			font-weight: 800;
			place-items: center;
		}

		.completion h1 {
			margin: 0 0 0.75rem;
			color: var(--navy);
			font-family: Georgia, "Times New Roman", serif;
			font-size: 2.5rem;
			line-height: 1.15;
		}

		.completion p {
			color: var(--muted);
			font-size: 1.15rem;
			line-height: 1.6;
		}

		.reference-box {
			margin: 1.75rem 0;
			padding: 1.5rem;
			border: 3px solid var(--green);
			border-radius: 8px;
			background: #f0f8f4;
			text-align: center;
		}

		.reference-box span {
			display: block;
			margin-bottom: 0.4rem;
			color: #355948;
			font-size: 1rem;
			font-weight: 800;
			text-transform: uppercase;
		}

		.reference-box strong {
			color: #174e36;
			font-family: "Courier New", monospace;
			font-size: 3.25rem;
			letter-spacing: 0.16em !important;
		}

		.arrival-note {
			margin: 1.25rem 0;
			padding: 1.25rem;
			border-top: 1px solid var(--line);
			border-bottom: 1px solid var(--line);
			color: var(--ink);
			font-size: 1.1rem;
			font-weight: 700;
			line-height: 1.55;
			text-align: center;
		}

		.chat-panel {
			margin: 1.5rem 0;
			border: 2px solid var(--line);
			border-radius: 8px;
			background: var(--surface);
			overflow: hidden;
		}

		.chat-topbar {
			display: flex;
			align-items: center;
			gap: 0.8rem;
			padding: 1rem 1.25rem;
			background: var(--navy);
			color: #ffffff;
		}

		.assistant-mark {
			display: grid;
			width: 2.8rem;
			height: 2.8rem;
			border-radius: 50%;
			background: #ffffff;
			color: var(--navy);
			font-size: 1.45rem;
			font-weight: 800;
			place-items: center;
		}

		.assistant-name {
			font-size: 1.1rem;
			font-weight: 800;
		}

		.assistant-status {
			font-size: 0.9rem;
			opacity: 0.85;
		}

		.conversation {
			display: flex;
			min-height: 10rem;
			padding: 1.25rem;
			flex-direction: column;
			gap: 1rem;
			background: #f7fafb;
		}

		.message {
			max-width: 82%;
			padding: 0.95rem 1rem;
			border-radius: 4px 8px 8px 8px;
			background: #e5f0f6;
			color: var(--ink);
			font-size: 1.08rem;
			line-height: 1.55;
		}

		.message.user {
			align-self: flex-end;
			border-radius: 8px 4px 8px 8px;
			background: var(--blue);
			color: #ffffff;
		}

		.chat-prompt {
			margin: 1.5rem 0 0.5rem;
			color: var(--ink);
			font-size: 1.05rem;
			font-weight: 800;
		}

		.site-footer {
			margin-top: 2.5rem;
			padding-top: 1rem;
			border-top: 1px solid var(--line);
			color: var(--muted);
			font-size: 0.9rem;
			line-height: 1.5;
			text-align: center;
		}

		@media (max-width: 640px) {
			.main .block-container {
				padding: 0.8rem 1rem 2rem;
			}

			.site-header {
				align-items: flex-start;
			}

			.brand {
				font-size: 1.05rem;
			}

			.brand-mark {
				width: 2.3rem;
				height: 2.3rem;
			}

			.help-number {
				font-size: 0.9rem;
			}

			.hero {
				padding: 2.5rem 0 1.5rem;
			}

			.hero h1, .flow-title h1, .completion h1 {
				font-size: 2rem;
			}

			.hero p, .flow-title p {
				font-size: 1.1rem;
			}

			.assurance-row {
				grid-template-columns: 1fr;
			}

			.assurance-item + .assurance-item {
				border-top: 1px solid var(--line);
				border-left: 0;
			}

			.reference-box strong {
				font-size: 2.5rem;
				letter-spacing: 0.1em !important;
			}

			.conversation {
				padding: 1rem;
			}

			.message {
				max-width: 92%;
			}
		}
	</style>
	""",
	unsafe_allow_html=True,
)

st.markdown(
	"""
	<header class="site-header">
		<div class="brand">
			<div class="brand-mark">+</div>
			<div>Medical Centre</div>
		</div>
		<div class="help-number">
			<span>Need help?</span>
			(555) 010-2020
		</div>
	</header>
	""",
	unsafe_allow_html=True,
)

if st.session_state.step == 0:
	st.markdown(
		"""
		<section class="hero">
			<h1>Spend less time at the front desk</h1>
			<p>Start your pre-arrival registration here, then follow the registration assistant one step at a time.</p>
		</section>
		""",
		unsafe_allow_html=True,
	)

	if st.button("Register Now"):
		st.session_state.step = 1
		st.rerun()

	st.markdown(
		"""
		<div class="assurance-row">
			<div class="assurance-item">About 3 minutes</div>
			<div class="assurance-item">Private and secure</div>
			<div class="assurance-item">Help is available</div>
		</div>
		""",
		unsafe_allow_html=True,
	)

elif st.session_state.step == 1:
	st.markdown(
		"""
		<div class="progress-label">Step 1 of 4: Start registration</div>
		<div class="progress-track"><div class="progress-fill" style="width: 25%"></div></div>
		<section class="flow-title">
			<h1>Pre-arrival registration</h1>
			<p>The assistant will guide you through each step.</p>
		</section>
		<div class="chat-panel">
			<div class="chat-topbar">
				<div class="assistant-mark">+</div>
				<div><div class="assistant-name">Registration Assistant</div><div class="assistant-status">Online now</div></div>
			</div>
			<div class="conversation">
				<div class="message">Hi there! How can I assist you today?</div>
			</div>
		</div>
		""",
		unsafe_allow_html=True,
	)

	trigger_phrase = st.text_input(
		"Your message",
		key="trigger_phrase_input",
		placeholder="I need to register",
	)

	if st.button("Send Message"):
		if trigger_phrase.strip().casefold() != "i need to register":
			st.warning('Please type "I need to register" to begin.')
		else:
			st.session_state.step = 2
			st.rerun()

elif st.session_state.step == 2:
	st.markdown(
		"""
		<div class="progress-label">Step 2 of 4: Patient details</div>
		<div class="progress-track"><div class="progress-fill" style="width: 50%"></div></div>
		<section class="flow-title">
			<h1>Tell us who you are</h1>
			<p>Enter the patient's name exactly as it appears on their identification.</p>
		</section>
		<div class="chat-panel">
			<div class="chat-topbar">
				<div class="assistant-mark">+</div>
				<div><div class="assistant-name">Registration Assistant</div><div class="assistant-status">Online now</div></div>
			</div>
			<div class="conversation">
				<div class="message user">I need to register</div>
				<div class="message">Of course. What is the patient's full name?</div>
			</div>
		</div>
		""",
		unsafe_allow_html=True,
	)

	full_name = st.text_input(
		"Patient's full name",
		key="patient_name_input",
		placeholder="For example: Maria Santos",
	)

	if st.button("Send Name"):
		if not full_name.strip():
			st.warning("Please enter the patient's full name.")
		else:
			st.session_state.patient_name = full_name.strip()
			st.session_state.step = 3
			st.rerun()

elif st.session_state.step == 3:
	patient_name = escape(st.session_state.patient_name)
	st.markdown(
		f"""
		<div class="progress-label">Step 3 of 4: Medical document</div>
		<div class="progress-track"><div class="progress-fill" style="width: 75%"></div></div>
		<section class="flow-title">
			<h1>Send your care document</h1>
			<p>Take a clear photo with your phone, or choose a PDF or Word document saved on your device.</p>
		</section>
		<div class="chat-panel">
			<div class="chat-topbar">
				<div class="assistant-mark">+</div>
				<div><div class="assistant-name">Registration Assistant</div><div class="assistant-status">Online now</div></div>
			</div>
			<div class="conversation">
				<div class="message user">{patient_name}</div>
				<div class="message">Thank you. Please upload your medical chit or authorization letter.</div>
			</div>
		</div>
		<p class="chat-prompt">Choose a clear image, PDF, or Word document:</p>
		""",
		unsafe_allow_html=True,
	)

	uploaded_file = st.file_uploader(
		"Upload your Medical Chit or Authorization Letter",
		type=["pdf", "docx", "png", "jpg", "jpeg"],
		help="Accepted files: PDF, DOCX, PNG, JPG, or JPEG.",
	)
	process_document = st.button("Process Document")

	if process_document:
		if uploaded_file is None:
			st.warning("Please take a photo or choose a document before continuing.")
		else:
			extracted_docx_text = None
			try:
				with st.spinner("Securely reading your document. Please wait..."):
					if uploaded_file.name.lower().endswith(".docx"):
						extracted_docx_text = read_docx_text(uploaded_file)
					time.sleep(1)
			except Exception:
				st.error(
					"We could not read this Word document. Please check the file and try again."
				)
			else:
				if extracted_docx_text == "":
					st.warning(
						"This Word document does not contain readable text. Please choose another file."
					)
				else:
					st.session_state.uploaded_file_name = uploaded_file.name
					st.session_state.extracted_docx_text = extracted_docx_text
					st.session_state.reference_code = str(
						100000 + secrets.randbelow(900000)
					)
					st.session_state.step = 4
					st.rerun()

elif st.session_state.step == 4:
	reference_code = st.session_state.reference_code
	st.markdown(
		f"""
		<div class="progress-label">Step 4 of 4: Complete</div>
		<div class="progress-track"><div class="progress-fill" style="width: 100%"></div></div>
		<section class="completion">
			<div class="success-mark">✓</div>
			<h1>Your registration is complete</h1>
			<p>The assistant has received your document and prepared your reference code.</p>
		</section>
		<div class="reference-box">
			<span>Your reference code</span>
			<strong>{reference_code}</strong>
		</div>
		<div class="arrival-note">Show this 6-digit code to the front desk when you arrive.</div>
		""",
		unsafe_allow_html=True,
	)

	if st.session_state.extracted_docx_text:
		st.success("Your Word document was read successfully.")
		with st.expander("Review text read from your Word document"):
			st.text_area(
				"Extracted document text",
				value=st.session_state.extracted_docx_text,
				height=240,
				disabled=True,
				label_visibility="collapsed",
			)

	render_copy_button(reference_code)

	if st.button("Start a New Registration"):
		for key in [
			"step",
			"reference_code",
			"patient_name",
			"patient_name_input",
			"trigger_phrase_input",
			"uploaded_file_name",
			"extracted_docx_text",
		]:
			st.session_state.pop(key, None)
		st.rerun()

st.markdown(
	"""
	<footer class="site-footer">
		Your information is handled securely. For help with registration, call (555) 010-2020.
	</footer>
	""",
	unsafe_allow_html=True,
)
